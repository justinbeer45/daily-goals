import docx

years = []


def _set_year_range():
    for i in range(2021, 2100):
        years.append(str(i))


months = ["January",
          "February",
          "March",
          "April",
          "May",
          "June",
          "July",
          "August",
          "September",
          "October",
          "November",
          "December"]


def get_text(filename):
    doc = docx.Document(filename)
    fulltext = []
    for paragraph in doc.paragraphs:
        fulltext.append(paragraph.text)
    return '\n\n'.join(fulltext)


def remove_pound_signs(filename):
    # create new document
    newdoc = docx.Document()
    # import existing document
    doc = docx.Document(filename)
    # for loop to iterate through original document and write non-pound sign characters to new document
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            string = run.text
            new_string = []
            for char in string:
                if char != "#":
                    new_string.append(char)
            newdoc.add_paragraph(new_string)

    # save document
    newdoc.save('NoPoundSigns.docx')


def remove_spaces(filename):
    # create new document
    newdoc = docx.Document()
    doc = docx.Document(filename)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            string = run.text
            string = string.strip()
            newdoc.add_paragraph(string)
    # save document
    newdoc.save('NoSpaces.docx')


def get_year_month(filename):
    year = ""
    month = ""
    _set_year_range()
    print(years)
    doc = docx.Document(filename)
    newdoc = docx.Document()
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            print(run.text.strip())
            if run.text.strip() in years:
                year = run.text.strip()
            if run.text.strip() in months:
                month = run.text.strip()
            #if year != "" and month != "":
                year_month = "{} {}, ".format(month.capitalize(), year)
                year_month_paragraph = newdoc.add_paragraph(year_month)
            if run.text.strip()[0] in ["1", "2", "3"] and run.text.strip() not in years:
                year_month_paragraph.add_run(run.text.strip())
    # save document
    newdoc.save('YearMonth.docx')
## TODO: August 2021 is missing the days for some reason